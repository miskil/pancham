import io
import logging
import unicodedata
from datetime import date

from docx import Document
from docx.shared import Pt
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from ..auth import get_current_user
from ..db import get_db
from ..models.plan import ProjectPlan
from ..models.proposal import Proposal
from ..models.village import Village

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/admin/export", tags=["admin-export"])


def _ascii_safe_filename(text: str) -> str:
    """Convert text to ASCII-safe filename by removing non-ASCII characters."""
    if not text:
        return "export"
    # Normalize and remove non-ASCII
    nfkd = unicodedata.normalize('NFKD', text)
    ascii_only = nfkd.encode('ascii', 'ignore').decode('ascii')
    # Remove special characters, keep alphanumeric and basic punctuation
    safe = ''.join(c if c.isalnum() or c in '-_' else '_' for c in ascii_only)
    return safe.strip('_') or "export"


def _docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.size = Pt(14 if level == 1 else 12)


def _row(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value or "—")


def _normalize_categories(milestone: dict) -> str:
    categories = milestone.get("categories") or milestone.get("category") or []
    if isinstance(categories, str):
        categories = [categories]
    cleaned = [str(cat).strip() for cat in categories if str(cat).strip()]
    return ", ".join(cleaned) or "—"


def _normalize_activities(milestone: dict) -> list[dict]:
    activities = milestone.get("activities")
    if isinstance(activities, list) and activities:
        return activities
    return [
        {
            "activity": milestone.get("activity") or milestone.get("details") or "",
            "poc": milestone.get("poc") or "",
            "amount": milestone.get("amount"),
            "notes": milestone.get("notes") or milestone.get("comment") or "",
        }
    ]


FOCUS_AREA_LABELS = {
    "HEALTH":              {"mr": "आरोग्य",             "en": "Health"},
    "EDUCATION":           {"mr": "शिक्षण",             "en": "Education"},
    "ENVIRONMENT":         {"mr": "पर्यावरण",           "en": "Environment"},
    "INCOME_GENERATION":   {"mr": "उत्पन्न निर्मिती",   "en": "Income Generation"},
    "WOMENS_EMPOWERMENT":  {"mr": "महिला सक्षमीकरण",   "en": "Women's Empowerment"},
}

# Each section: (js_key, {mr, en}, [ (field_key, {mr, en}), ... ])
VILLAGE_PROFILE_SECTIONS = [
    ("basic", {"mr": "गावाची मूलभूत माहिती", "en": "Basic Village Information"}, [
        ("population",         {"mr": "लोकसंख्या",                               "en": "Population"}),
        ("total_families",     {"mr": "एकूण कुटुंबांची संख्या",                 "en": "Total Families"}),
        ("area",               {"mr": "गावाचे क्षेत्रफळ",                        "en": "Village Area"}),
        ("sc_st_obc_ratio",    {"mr": "SC/ST/OBC व इतर लोकसंख्येचे प्रमाण",     "en": "SC/ST/OBC & Other Population Ratio"}),
        ("literacy_rate",      {"mr": "साक्षरता दर",                             "en": "Literacy Rate"}),
        ("migration_families", {"mr": "स्थलांतर करणाऱ्या कुटुंबांची संख्या",    "en": "Migrating Families"}),
    ]),
    ("poverty", {"mr": "गरिबी व असुरक्षितता", "en": "Poverty & Vulnerability"}, [
        ("bpl_families",          {"mr": "गरीबी रेषेखालील कुटुंबांचे प्रमाण", "en": "Below Poverty Line (BPL) Families %"}),
        ("landless_families",     {"mr": "भूमिहीन कुटुंबे",                    "en": "Landless Families"}),
        ("wage_dependent",        {"mr": "मजुरीवर अवलंबून कुटुंबे",             "en": "Wage-Dependent Families"}),
        ("malnourished_children", {"mr": "कुपोषित बालकांची संख्या",            "en": "Malnourished Children"}),
        ("women_headed",          {"mr": "महिला प्रमुख कुटुंबे",                "en": "Women-Headed Families"}),
        ("disabled",              {"mr": "दिव्यांग व्यक्ती",                    "en": "Persons with Disability"}),
        ("elderly_alone",         {"mr": "वृद्ध व एकटे राहणारे नागरिक",        "en": "Elderly Living Alone"}),
    ]),
    ("water", {"mr": "पाणी व जलसुरक्षा", "en": "Water & Water Security"}, [
        ("drinking_sources",    {"mr": "पिण्याच्या पाण्याचे स्रोत",   "en": "Drinking Water Sources"}),
        ("scarcity_months",     {"mr": "वर्षातील पाणीटंचाईचे महिने",  "en": "Water Scarcity Months/Year"}),
        ("groundwater_level",   {"mr": "भूजल पातळी",                   "en": "Groundwater Level"}),
        ("irrigated_area",      {"mr": "सिंचनाखालील क्षेत्र",          "en": "Irrigated Area"}),
        ("tanker_dependency",   {"mr": "टँकरवर अवलंबित्व",             "en": "Tanker Dependency"}),
        ("conservation_status", {"mr": "जलसंधारणाची स्थिती",           "en": "Water Conservation Status"}),
    ]),
    ("agriculture", {"mr": "शेती व उपजीविका", "en": "Agriculture & Livelihoods"}, [
        ("main_crops",          {"mr": "प्रमुख पिके",                   "en": "Main Crops"}),
        ("irrigated_rainfed",   {"mr": "सिंचित व जिरायती क्षेत्र",      "en": "Irrigated & Rainfed Area"}),
        ("crop_productivity",   {"mr": "पीक उत्पादकता",                 "en": "Crop Productivity"}),
        ("livestock",           {"mr": "पशुधन संख्या",                  "en": "Livestock Count"}),
        ("fpo_exists",          {"mr": "FPO अस्तित्वात आहे का?",        "en": "FPO Exists?"}),
        ("non_farm_employment", {"mr": "कृषीबाह्य रोजगाराच्या संधी",    "en": "Non-Farm Employment Opportunities"}),
        ("avg_income",          {"mr": "कुटुंबांचे सरासरी उत्पन्न",     "en": "Average Family Income"}),
    ]),
    ("infrastructure", {"mr": "पायाभूत सुविधा", "en": "Infrastructure"}, [
        ("roads",            {"mr": "रस्ते संपर्क",      "en": "Road Connectivity"}),
        ("electricity",      {"mr": "वीज उपलब्धता",      "en": "Electricity Availability"}),
        ("mobile_network",   {"mr": "मोबाईल नेटवर्क",    "en": "Mobile Network"}),
        ("internet",         {"mr": "इंटरनेट सुविधा",     "en": "Internet Facility"}),
        ("public_transport", {"mr": "सार्वजनिक वाहतूक",  "en": "Public Transport"}),
        ("housing",          {"mr": "घरांची स्थिती",      "en": "Housing Condition"}),
    ]),
    ("education", {"mr": "शिक्षण", "en": "Education"}, [
        ("schools",           {"mr": "प्राथमिक व माध्यमिक शाळांची उपलब्धता", "en": "Primary & Secondary Schools"}),
        ("attendance",        {"mr": "शाळेत उपस्थिती",                        "en": "School Attendance"}),
        ("out_of_school",     {"mr": "शाळाबाह्य मुलांचे प्रमाण",              "en": "Out-of-School Children"}),
        ("girls_education",   {"mr": "मुलींच्या शिक्षणाची स्थिती",            "en": "Girls' Education Status"}),
        ("digital_education", {"mr": "डिजिटल शिक्षण सुविधा",                  "en": "Digital Education Facilities"}),
    ]),
    ("health", {"mr": "आरोग्य व पोषण", "en": "Health & Nutrition"}, [
        ("phc_distance",          {"mr": "PHC पर्यंतचे अंतर",               "en": "Distance to PHC"}),
        ("asha_anganwadi",        {"mr": "आशा व अंगणवाडी सेविका उपलब्धता", "en": "ASHA & Anganwadi Workers"}),
        ("maternal_child_health", {"mr": "माता व बाल आरोग्य स्थिती",        "en": "Maternal & Child Health"}),
        ("anemia",                {"mr": "रक्तक्षय (Anemia)",                "en": "Anemia"}),
        ("vaccination",           {"mr": "लसीकरणाचे प्रमाण",                "en": "Vaccination Coverage"}),
    ]),
    ("governance", {"mr": "स्थानिक स्वराज्य संस्था व प्रशासन", "en": "Local Governance"}, [
        ("gramsabha_count",        {"mr": "ग्रामसभांची संख्या व नियमितता", "en": "Gram Sabha Count & Regularity"}),
        ("gramsabha_attendance",   {"mr": "ग्रामसभेतील उपस्थिती",          "en": "Gram Sabha Attendance"}),
        ("women_participation",    {"mr": "महिलांचा सहभाग",                 "en": "Women's Participation"}),
        ("gpdp_quality",           {"mr": "GPDP ची गुणवत्ता",               "en": "GPDP Quality"}),
        ("committees",             {"mr": "समित्यांचे कार्य",                "en": "Committee Functioning"}),
        ("financial_transparency", {"mr": "आर्थिक पारदर्शकता",             "en": "Financial Transparency"}),
        ("scheme_implementation",  {"mr": "सरकारी योजनांची अंमलबजावणी",    "en": "Govt Scheme Implementation"}),
    ]),
    ("community", {"mr": "समुदायाची तयारी", "en": "Community Readiness"}, [
        ("shramdan",              {"mr": "श्रमदानाची तयारी",           "en": "Readiness for Shramdan"}),
        ("financial_contribution",{"mr": "आर्थिक योगदानाची तयारी",    "en": "Financial Contribution Readiness"}),
        ("shg_active",            {"mr": "सक्रिय बचत गट",             "en": "Active Self-Help Groups (SHG)"}),
        ("youth_groups",          {"mr": "युवक मंडळे",                 "en": "Youth Groups"}),
        ("farmer_groups",         {"mr": "शेतकरी गट",                  "en": "Farmer Groups"}),
        ("local_leadership",      {"mr": "स्थानिक नेतृत्व",             "en": "Local Leadership"}),
        ("collective_history",    {"mr": "सामूहिक कामांचा इतिहास",    "en": "History of Collective Work"}),
    ]),
    ("projects", {"mr": "चालू सरकारी व इतर संस्थांचे प्रकल्प", "en": "Ongoing Projects"}, [
        ("water_conservation",    {"mr": "जलसंधारण प्रकल्प",          "en": "Water Conservation Projects"}),
        ("csr_projects",          {"mr": "CSR प्रकल्प",                "en": "CSR Projects"}),
        ("ngo_projects",          {"mr": "इतर NGO प्रकल्प",           "en": "Other NGO Projects"}),
        ("agriculture_programs",  {"mr": "कृषी विकास कार्यक्रम",       "en": "Agriculture Development Programs"}),
        ("infrastructure_schemes",{"mr": "पायाभूत सुविधा योजना",       "en": "Infrastructure Schemes"}),
    ]),
    ("environment", {"mr": "पर्यावरण व हवामान जोखीम", "en": "Environment & Climate Risk"}, [
        ("drought_prone",  {"mr": "दुष्काळग्रस्तता",         "en": "Drought Proneness"}),
        ("flood_risk",     {"mr": "पूर जोखीम",               "en": "Flood Risk"}),
        ("soil_erosion",   {"mr": "मृदा धूप",                "en": "Soil Erosion"}),
        ("forest_area",    {"mr": "जंगल क्षेत्र",             "en": "Forest Area"}),
        ("climate_impact", {"mr": "हवामान बदलाचा परिणाम",    "en": "Climate Change Impact"}),
    ]),
    ("tribal", {"mr": "आदिवासी व PESA स्थिती", "en": "Tribal & PESA Status"}, [
        ("scheduled_area",    {"mr": "गाव अनुसूचित क्षेत्रात आहे का?", "en": "Village in Scheduled Area?"}),
        ("pesa_applicable",   {"mr": "PESA लागू आहे का?",              "en": "PESA Applicable?"}),
        ("forest_rights",     {"mr": "वनहक्क दावे",                     "en": "Forest Rights Claims"}),
        ("forest_dependency", {"mr": "वनसंपत्तीवरील अवलंबित्व",         "en": "Dependence on Forest Resources"}),
    ]),
]

EXPORT_LABELS = {
    "mr": {
        "proposal":        "प्रस्ताव",
        "village":         "गाव",
        "district":        "जिल्हा",
        "taluka":          "तालुका",
        "status":          "स्थिती",
        "submitted":       "सादर केले",
        "not_submitted":   "सादर केले नाही",
        "ngo_partner":     "NGO भागीदार",
        "ngo_name":        "NGO चे नाव",
        "fcra_number":     "FCRA क्रमांक",
        "fcra_expiry":     "FCRA समाप्ती तारीख",
        "ngo_lead":        "NGO प्रमुख",
        "ngo_lead_phone":  "NGO प्रमुख फोन",
        "whatsapp":        "व्हॉट्सॲप",
        "bank_account":    "बँक खाते क्रमांक",
        "ifsc":            "IFSC कोड",
        "village_lead":    "गाव प्रतिनिधी",
        "name":            "नाव",
        "phone":           "फोन",
        "focus_areas":     "फोकस क्षेत्रे",
        "geo_desc":        "भौगोलिक व सामाजिक माहिती",
        "community_ctx":   "गावाची गरज",
        "key_activities":  "नियोजित प्रमुख उपक्रम",
        "reviewer_notes":  "पुनरावलोकन नोंदी",
        "village_profile": "गावाची सविस्तर माहिती",
    },
    "en": {
        "proposal":        "Proposal",
        "village":         "Village",
        "district":        "District",
        "taluka":          "Taluka",
        "status":          "Status",
        "submitted":       "Submitted",
        "not_submitted":   "Not submitted",
        "ngo_partner":     "NGO Partner",
        "ngo_name":        "NGO Name",
        "fcra_number":     "FCRA Number",
        "fcra_expiry":     "FCRA Expiry",
        "ngo_lead":        "NGO Lead",
        "ngo_lead_phone":  "NGO Lead Phone",
        "whatsapp":        "WhatsApp",
        "bank_account":    "Bank Account Number",
        "ifsc":            "IFSC Code",
        "village_lead":    "Village Lead",
        "name":            "Name",
        "phone":           "Phone",
        "focus_areas":     "Focus Areas",
        "geo_desc":        "Geographic & Social Description",
        "community_ctx":   "Village Needs / Community Context",
        "key_activities":  "Key Activities Planned",
        "reviewer_notes":  "Reviewer Notes",
        "village_profile": "Village Profile",
    },
}


# ── Proposal export ──────────────────────────────────────────────────────────

@router.post("/proposals/{proposal_id}")
async def export_proposal(
    proposal_id: str,
    lang: str = Query("mr", pattern="^(mr|en)$"),
    db: AsyncSession = Depends(get_db),
    user: dict = Depends(get_current_user),
):
    try:
        result = await db.execute(
            select(Proposal).where(Proposal.id == proposal_id)
        )
        proposal = result.scalar_one_or_none()
        if not proposal:
            raise HTTPException(status_code=404, detail="Proposal not found")

        if user["role"] not in ["ADMIN", "VILLAGE"]:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        if user["role"] == "VILLAGE" and user.get("village_id") != proposal.village_id:
            raise HTTPException(status_code=403, detail="Insufficient permissions")

        village = await db.get(Village, proposal.village_id)
        L = EXPORT_LABELS[lang]

        doc = Document()
        doc.core_properties.author = "Pancham"

        village_name = village.name if village else proposal.village_id
        _heading(doc, f"{L['proposal']} — {village_name}")
        _row(doc, L["village"], village_name)
        _row(doc, L["district"], village.district if village else "")
        _row(doc, L["taluka"], village.taluka if village else "")
        _row(doc, L["status"], proposal.status)
        _row(doc, L["submitted"], str(proposal.submitted_at.date()) if proposal.submitted_at else L["not_submitted"])
        doc.add_paragraph()

        _heading(doc, L["ngo_partner"], level=2)
        _row(doc, L["ngo_name"], village.ngo_name or "—")
        _row(doc, L["fcra_number"], village.fcra_number or "—")
        _row(doc, L["fcra_expiry"], str(village.fcra_expiry_date) if village.fcra_expiry_date else "—")
        _row(doc, L["ngo_lead"], village.ngo_contact_name or "—")
        _row(doc, L["ngo_lead_phone"], village.ngo_contact_phone or "—")
        if village.ngo_whatsapp_phone:
            _row(doc, L["whatsapp"], village.ngo_whatsapp_phone)
        _row(doc, L["bank_account"], village.bank_account_number or "—")
        _row(doc, L["ifsc"], village.ifsc_code or "—")
        doc.add_paragraph()

        _heading(doc, L["village_lead"], level=2)
        _row(doc, L["name"], village.village_lead_name or "—")
        _row(doc, L["phone"], village.village_lead_phone or "—")
        doc.add_paragraph()

        _heading(doc, L["focus_areas"], level=2)
        raw_areas = [item.strip() for item in (proposal.focus_area or "").split(",") if item.strip()]
        labeled_areas = ", ".join(FOCUS_AREA_LABELS.get(a, {}).get(lang, a) for a in raw_areas) or "—"
        doc.add_paragraph(labeled_areas)
        doc.add_paragraph()

        _heading(doc, L["geo_desc"], level=2)
        doc.add_paragraph(proposal.description or "—")

        _heading(doc, L["community_ctx"], level=2)
        doc.add_paragraph(proposal.community_context or "—")

        _heading(doc, L["key_activities"], level=2)
        doc.add_paragraph(proposal.key_activities or "—")

        _heading(doc, L["village_profile"], level=1)
        profile = village.village_profile or {}
        for js_key, section_labels, fields in VILLAGE_PROFILE_SECTIONS:
            section_data = profile.get(js_key, {})
            _heading(doc, section_labels[lang], level=2)
            for field_key, field_labels in fields:
                val = section_data.get(field_key)
                _row(doc, field_labels[lang], str(val) if val else "—")
            doc.add_paragraph()

        if proposal.reviewer_notes:
            _heading(doc, L["reviewer_notes"], level=2)
            doc.add_paragraph(proposal.reviewer_notes)

        village_slug = _ascii_safe_filename(village_name)
        filename = f"{village_slug}-Proposal-{date.today()}.docx"
        content = _docx_bytes(doc)
        headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export proposal failed for {proposal_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ── Plan export ───────────────────────────────────────────────────────────────

@router.post("/plans/{plan_id}")
async def export_plan(
    plan_id: str,
    db: AsyncSession = Depends(get_db),
    user: dict = Depends(get_current_user),
):
    try:
        result = await db.execute(
            select(ProjectPlan).where(ProjectPlan.id == plan_id)
        )
        plan = result.scalar_one_or_none()
        if not plan:
            raise HTTPException(status_code=404, detail="Plan not found")

        if user["role"] not in ["ADMIN", "VILLAGE"]:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        if user["role"] == "VILLAGE" and user.get("village_id") != plan.village_id:
            raise HTTPException(status_code=403, detail="Insufficient permissions")

        village = await db.get(Village, plan.village_id)

        doc = Document()
        doc.core_properties.author = "Pancham"

        _heading(doc, f"Project Plan — {village.name if village else plan.village_id}")
        _row(doc, "Village", village.name if village else "")
        _row(doc, "District", village.district if village else "")
        _row(doc, "Version", plan.version_type)
        _row(doc, "Status", plan.status)
        _row(doc, "Start Date", str(plan.start_date) if plan.start_date else "—")
        _row(doc, "End Date", str(plan.end_date) if plan.end_date else "—")
        if plan.frozen_at:
            _row(doc, "Frozen At", str(plan.frozen_at.date()))
        doc.add_paragraph()

        _heading(doc, "Village Lead", level=2)
        _row(doc, "Name", village.village_lead_name or "—")
        _row(doc, "Phone", village.village_lead_phone or "—")
        doc.add_paragraph()

        _heading(doc, "NGO Partner", level=2)
        _row(doc, "NGO Name", village.ngo_name or "—")
        _row(doc, "Contact Name", village.ngo_contact_name or "—")
        _row(doc, "Phone", village.ngo_contact_phone or "—")
        if village.ngo_whatsapp_phone:
            _row(doc, "WhatsApp", village.ngo_whatsapp_phone)
        doc.add_paragraph()

        plan_data = plan.plan_data or {}
        for year_key in ["1", "2", "3"]:
            milestones = plan_data.get(year_key, [])
            if not milestones:
                continue
            _heading(doc, f"Year {year_key}", level=2)
            for idx, milestone in enumerate(milestones, start=1):
                milestone_title = milestone.get("milestone") or milestone.get("title") or f"Milestone {idx}"
                _heading(doc, milestone_title, level=3)
                _row(doc, "Categories", _normalize_categories(milestone))
                _row(doc, "Impact", milestone.get("impact") or milestone.get("impact_box") or "")

                table = doc.add_table(rows=1, cols=4)
                table.style = "Light List Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "Activity"
                hdr[1].text = "POC"
                hdr[2].text = "Notes"
                hdr[3].text = "Amount"

                for activity in _normalize_activities(milestone):
                    cells = table.add_row().cells
                    cells[0].text = activity.get("activity", "") or "—"
                    cells[1].text = activity.get("poc", "") or "—"
                    cells[2].text = activity.get("notes", "") or "—"
                    amount = activity.get("amount")
                    cells[3].text = str(amount) if amount is not None else "—"
                doc.add_paragraph()

        village_slug = _ascii_safe_filename(village.name if village else plan.village_id)
        filename = f"{village_slug}-Plan-{date.today()}.docx"
        content = _docx_bytes(doc)
        headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Export plan failed for {plan_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")
